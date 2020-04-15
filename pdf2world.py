from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from docx import Document
import io


def remove_control_char(pdf_content):
    """
    移除控制字符，换行符、制表符、转义符等）python-docx是不支持控制字符写入的
    :param pdf_content: PDF 文件内容
    :return:  返回去除控制字符的内容
    """

    # 使用str的translate方法，将ASCII码在32以下的都移除 http://ascii.911cha.com/
    return pdf_content.translate(dict.fromkeys(range(32)))


def read_from_pdf(pdf_file):
    """
    读取PDF文件内容，并做处理
    :param pdf_file: PDF 文件
    :return: pdf文件内容
    """

    # 二进制读取pdf文件内的内容
    with open(pdf_file, 'rb') as file:
        # 创建PDF资源管理器，管理共享资源
        resource_manage = PDFResourceManager()
        return_str = io.StringIO()

        # 创建一个PDF设备对象
        lap_params = LAParams()

        # 内容转换
        device = TextConverter(
            resource_manage,
            return_str,
            laparams = lap_params
        )
        process_pdf(resource_manage, device, file)
        device.close()
        # 获取转换后的pdf文件内容
        pdf_content = return_str.getvalue()
#        print(pdf_content)
        return pdf_content


def write_to_world(world_file, pdf_content):
    """
    将处理过的pdf文件内容写入到 world 中进行保存
    :param world_file: world 文件
    :param pdf_content: PDF 文件内容
    :return: pdf文件内容
    """
    docx = Document()
    # 之前读取的是字符串，所以需要split将其分隔为每一行
    for row in pdf_content.split("\n"):
        paragraph = docx.add_paragraph()
        paragraph.add_run(remove_control_char(row))
    docx.save(world_file)

def pdf_to_world(world_file, pdf_file):
    pdf_content = read_from_pdf(pdf_file)
    write_to_world(world_file, pdf_content)


if __name__ == '__main__':
    pdf_to_world("./1.docx", "./test.pdf")


